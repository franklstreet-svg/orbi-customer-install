"""
doc_convert — take any source document, clean it up with the LLM, and write
it back out in the target format the owner wants.

Pipeline:
    source file (any supported type)
      → extract_text()  -- get plain text from PDF/docx/xlsx/csv/txt/md/html
      → llm_clean()     -- optional: LLM fixes typos / normalizes formatting
      → write_<fmt>()   -- emit the chosen target format

Targets supported in v1:
    pdf   — reportlab (simple text-based PDF, headings preserved)
    docx  — python-docx (real Word file with paragraph styles)
    txt   — plain UTF-8
    md    — markdown (best-effort; LLM produces markdown)
    xlsx  — openpyxl (only meaningful when source has tabular data)
    csv   — only meaningful when source has tabular data

The cleaned file is saved into the owner's workspace folder so it shows up
in the Files tab AND can be downloaded via the existing file-fetch path.
"""

from __future__ import annotations

import csv
import io
import logging
import re
import time
from pathlib import Path

log = logging.getLogger("orbi.doc_convert")

SUPPORTED_TARGETS = ("pdf", "docx", "txt", "md", "xlsx", "csv")
MAX_INPUT_CHARS = 30_000      # LLM cleanup capped here
MAX_OUTPUT_CHARS = 60_000     # safety upper bound on output


# ---------------------------------------------------------------------------
# Text extraction (delegates to workspace module since it already supports
# the same formats)
# ---------------------------------------------------------------------------


def extract_text(path: Path) -> str:
    """Read any supported file type and return its text content as a string."""
    from modules.workspace import _extract_text
    return _extract_text(path) or ""


def detect_kind(path: Path) -> str:
    """Categorize a source for routing decisions ('table' vs 'prose')."""
    suffix = path.suffix.lower()
    if suffix in (".csv", ".xlsx", ".xls"):
        return "table"
    return "prose"


# ---------------------------------------------------------------------------
# LLM cleanup
# ---------------------------------------------------------------------------


def llm_clean(config: dict, text: str, target: str, kind: str = "prose",
              hint: str = "") -> str:
    """Send text to the LLM with a 'clean this up' instruction. Returns the
    cleaned text. If the LLM is unreachable, returns the input unchanged so
    conversion still works (just without the cleanup polish)."""
    import llm_client
    if not text or not text.strip():
        return text
    if len(text) > MAX_INPUT_CHARS:
        text = text[:MAX_INPUT_CHARS]
        text += "\n\n[Note: input was truncated to fit the cleanup window.]"

    if kind == "table":
        system = _system_for_table(target, hint)
    else:
        system = _system_for_prose(target, hint)
    user_msg = ("Clean this content up per the rules above. "
                "Output ONLY the cleaned content, no preamble or commentary.\n\n"
                "---\n\n" + text)
    try:
        resp = llm_client.generate(config, system, [{"role": "user", "content": user_msg}])
        out = (resp.text or "").strip()
        if out:
            # Strip common LLM preambles
            out = re.sub(r"^(here(?:'s| is) (?:the )?(?:cleaned|cleaned-up|polished)[^\n:]*:?\s*\n)",
                         "", out, flags=re.IGNORECASE)
            return out[:MAX_OUTPUT_CHARS]
    except Exception as e:
        log.warning(f"llm cleanup failed, returning original: {e}")
    return text


def _system_for_prose(target: str, hint: str) -> str:
    target_note = {
        "pdf":  "The result will be rendered as a PDF. Use clear paragraph breaks and section headings (markdown # / ## are fine).",
        "docx": "The result will become a Word document. Use markdown # / ## for headings; they map to Word heading styles.",
        "md":   "Output markdown. Use # / ## for headings and standard markdown formatting.",
        "txt":  "Output plain text only. No markdown symbols. Use blank lines between paragraphs.",
        "csv":  "Output CSV: header row + data rows. Comma-separated. Quote fields with commas.",
        "xlsx": "Output as a tab-separated table: header row then data rows, each row on its own line, columns separated by single tabs.",
    }.get(target, "Plain text output.")
    extra = f"\nADDITIONAL OWNER INSTRUCTIONS: {hint}" if hint and hint.strip() else ""
    return f"""You are a document-cleanup assistant. Your job:

1. Fix obvious typos and grammar mistakes — but DO NOT change the meaning.
2. Normalize formatting: consistent paragraph breaks, consistent capitalization
   in headings, no extra blank lines, no weird whitespace.
3. If the original has clear sections, give them proper headings.
4. If the original is structured info (a list, a table), preserve the structure.
5. NEVER add new information the original doesn't have. NEVER summarize unless
   explicitly asked. Preserve every fact and every name.
6. If the original is already clean, return it nearly verbatim.

TARGET FORMAT: {target.upper()}
{target_note}

DO NOT include any preamble like "Here is the cleaned version:" — output only
the cleaned content itself. No backticks around the result.{extra}"""


def _system_for_table(target: str, hint: str) -> str:
    extra = f"\nADDITIONAL OWNER INSTRUCTIONS: {hint}" if hint and hint.strip() else ""
    if target == "pdf":
        format_rule = "Output the cleaned table as tab-separated rows: header row, then data rows."
    elif target == "csv":
        format_rule = "Output as CSV — comma separated, header row first, quote fields containing commas."
    elif target == "xlsx":
        format_rule = "Output as tab-separated rows — header row first, then data rows."
    elif target == "md":
        format_rule = "Output a markdown table (pipe-separated)."
    else:
        format_rule = "Output as tab-separated rows — header row then data rows."
    return f"""You are a spreadsheet-cleanup assistant. Your job:

1. Standardize column headers (consistent capitalization, no awkward names).
2. Remove obvious duplicate rows.
3. Fix obvious data-type issues (stray quotes, inconsistent date formats).
4. DO NOT change the actual data values unless they're clearly errors.
5. DO NOT add new rows or new columns. DO NOT summarize.

TARGET FORMAT: {target.upper()}
{format_rule}

Output ONLY the table content — no preamble, no explanation.{extra}"""


# ---------------------------------------------------------------------------
# Writers — one per target format
# ---------------------------------------------------------------------------


def write_txt(text: str, out_path: Path) -> Path:
    out_path.write_text(text, encoding="utf-8")
    return out_path


def write_md(text: str, out_path: Path) -> Path:
    out_path.write_text(text, encoding="utf-8")
    return out_path


def write_pdf(text: str, out_path: Path, title: str = "Orbi cleaned document") -> Path:
    """Render the cleaned text as a clean PDF. Recognizes markdown-style
    headings (# / ## / ###) and turns them into PDF headings. Everything
    else becomes wrapped paragraph text."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     PageBreak)
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT

    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"],
                          fontSize=11, leading=15, alignment=TA_LEFT,
                          spaceAfter=8)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                        fontSize=18, leading=22, spaceBefore=12, spaceAfter=10)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"],
                        fontSize=14, leading=18, spaceBefore=10, spaceAfter=8)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"],
                        fontSize=12, leading=16, spaceBefore=8, spaceAfter=6)

    doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
        title=title,
    )

    story = []
    for line in text.splitlines():
        line = line.rstrip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        if line.startswith("### "):
            story.append(Paragraph(_pdf_escape(line[4:]), h3))
        elif line.startswith("## "):
            story.append(Paragraph(_pdf_escape(line[3:]), h2))
        elif line.startswith("# "):
            story.append(Paragraph(_pdf_escape(line[2:]), h1))
        elif line.startswith("---"):
            story.append(Spacer(1, 12))
        else:
            story.append(Paragraph(_pdf_escape(line), body))
    if not story:
        story = [Paragraph("(empty)", body)]
    doc.build(story)
    return out_path


def _pdf_escape(s: str) -> str:
    """Escape characters that mean something to reportlab's mini-HTML parser."""
    return (s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;"))


def write_docx(text: str, out_path: Path, title: str = "Orbi cleaned document",
               *, header_text: str = "", footer_text: str = "",
               page_numbers: bool = False, business_name: str = "") -> Path:
    """Write the cleaned text as a real Word file.

    Markdown features supported:
      # / ## / ### / ####          → Word heading styles
      | a | b | c |                 → real Word tables (consecutive pipe lines)
      | - | - | - |                 → table-header separator (skipped)
      blank line                    → paragraph break
      ---                           → horizontal-rule paragraph

    Optional document furniture:
      header_text / footer_text     → page header / footer band
      page_numbers                  → adds "Page X of Y" in the footer
      business_name                 → defaults the header if set + header_text empty
    """
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document()
    d.core_properties.title = title

    # Header + footer band on the default section
    if header_text or business_name:
        section = d.sections[0]
        hdr = section.header.paragraphs[0]
        hdr.text = header_text or business_name
    if footer_text or page_numbers:
        section = d.sections[0]
        ftr = section.footer.paragraphs[0]
        if footer_text:
            ftr.text = footer_text
        if page_numbers:
            # Add a "Page X of Y" field. Word field-code XML trick.
            run = ftr.add_run("  Page ")
            _add_field_code(run, "PAGE")
            ftr.add_run(" of ")
            run2 = ftr.add_run()
            _add_field_code(run2, "NUMPAGES")

    # Walk the text. Detect runs of pipe-table lines and emit a real table.
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # Markdown table block?
        if _is_pipe_table_row(line):
            block = [line]
            j = i + 1
            while j < len(lines) and _is_pipe_table_row(lines[j].rstrip()):
                block.append(lines[j].rstrip())
                j += 1
            _emit_docx_table(d, block)
            i = j
            continue
        # Plain content
        if not line:
            d.add_paragraph("")
        elif line.startswith("#### "):
            d.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            d.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            d.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            d.add_heading(line[2:], level=1)
        elif line.startswith("---") and set(line) <= {"-", " "}:
            p = d.add_paragraph()
            p.add_run("_" * 60)
        elif line.startswith("- ") or line.startswith("* "):
            d.add_paragraph(line[2:], style="List Bullet")
        else:
            d.add_paragraph(line)
        i += 1

    d.save(str(out_path))
    return out_path


def _is_pipe_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2 and s.endswith("|")


def _emit_docx_table(doc, block: list[str]) -> None:
    """Convert a markdown pipe-table block into a real Word table."""
    rows = []
    for ln in block:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        # Skip pure-separator rows (|---|---|)
        if cells and all(set(c) <= set("-: ") for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < n_cols:
                table.cell(r_idx, c_idx).text = cell_text


def _add_field_code(run, code: str) -> None:
    """Insert a Word field code (e.g. PAGE, NUMPAGES) into a docx run."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = code
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def write_xlsx(text: str, out_path: Path, title: str = "Orbi cleaned data") -> Path:
    """Take cleaned table data and write XLSX with formulas, multi-sheet
    support, and basic styling.

    Multi-sheet syntax — the LLM (or owner) can mark sheet breaks in the
    incoming text with a line like:
        ## Sheet: Q1 Sales
        Header1\tHeader2\tHeader3
        ...rows...
        ## Sheet: Q2 Sales
        ...rows...

    Formula syntax — any cell value starting with '=' is written as a
    formula, not literal text. e.g. =SUM(B2:B10), =AVERAGE(C:C), etc.

    Styling: header row is bold + light gray fill, columns autosized.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    # Remove the default blank sheet — we'll add our own
    default = wb.active
    wb.remove(default)

    sheets = _split_sheets(text)
    if not sheets:
        sheets = [("Sheet1", [])]

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4F8CFF", end_color="4F8CFF",
                              fill_type="solid")
    header_align = Alignment(horizontal="left", vertical="center")

    for sheet_name, rows in sheets:
        ws = wb.create_sheet(title=(sheet_name or "Sheet")[:31])  # 31 = Excel max
        if not rows:
            continue
        # Append rows. If a cell starts with "=" treat as formula.
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, raw in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx)
                val = (raw or "").strip()
                if val.startswith("="):
                    # Real formula
                    cell.value = val
                else:
                    # Try to coerce numeric strings to numbers so Excel
                    # treats them right (sums, charts, etc.)
                    cell.value = _coerce_cell(val)
        # Style the header row
        for c_idx in range(1, len(rows[0]) + 1):
            cell = ws.cell(row=1, column=c_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
        # Autosize columns (best-effort — actual width is approximate)
        for c_idx in range(1, len(rows[0]) + 1):
            max_len = 8
            for r in rows[:50]:  # sample first 50 rows for sizing
                if c_idx <= len(r):
                    max_len = max(max_len, min(40, len(str(r[c_idx - 1]))))
            ws.column_dimensions[get_column_letter(c_idx)].width = max_len + 2
        # Freeze header row
        ws.freeze_panes = "A2"

    wb.properties.title = title
    wb.save(str(out_path))
    return out_path


def _coerce_cell(val: str):
    """Convert numeric-looking strings into actual numbers so Excel treats
    them as numbers (and they sum correctly, chart correctly, etc.)."""
    if not val:
        return val
    # Strip $ and , for currency-style values
    stripped = val.replace("$", "").replace(",", "").strip()
    try:
        if "." in stripped:
            return float(stripped)
        return int(stripped)
    except ValueError:
        return val


def _split_sheets(text: str) -> list[tuple[str, list[list[str]]]]:
    """Detect ## Sheet: <name> markers and split into multiple sheets.
    If no markers, returns a single-sheet list."""
    lines = text.splitlines()
    sheets: list[tuple[str, list[str]]] = []
    current_name = "Sheet1"
    current_lines: list[str] = []
    for ln in lines:
        m_match = ln.strip()
        if m_match.lower().startswith("## sheet:") or m_match.lower().startswith("# sheet:"):
            # Flush current
            if current_lines:
                sheets.append((current_name, current_lines))
            current_name = m_match.split(":", 1)[1].strip() or f"Sheet{len(sheets)+1}"
            current_lines = []
            continue
        current_lines.append(ln)
    if current_lines:
        sheets.append((current_name, current_lines))
    # Parse each sheet's lines into rows via existing _parse_table
    return [(name, _parse_table("\n".join(block))) for name, block in sheets]


def write_csv(text: str, out_path: Path) -> Path:
    rows = _parse_table(text)
    with out_path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        for row in rows:
            writer.writerow(row)
    return out_path


def _parse_table(text: str) -> list[list[str]]:
    """Detect whether the LLM output is tab- or comma-separated and parse it."""
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return []
    # If the first line has tabs OR pipes, treat as tab/pipe-separated
    first = lines[0]
    if "\t" in first:
        return [ln.split("\t") for ln in lines]
    if first.count("|") >= 2:
        # Markdown table — strip leading/trailing pipes, drop separator row
        out = []
        for ln in lines:
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            # Skip pure-separator rows like |---|---|
            if cells and all(set(c) <= set("-: ") for c in cells):
                continue
            out.append(cells)
        return out
    # Fallback: CSV
    buf = io.StringIO(text)
    return [row for row in csv.reader(buf)]


# ---------------------------------------------------------------------------
# Public end-to-end pipeline
# ---------------------------------------------------------------------------


def convert(config: dict, source_path: Path, target: str,
            out_dir: Path, hint: str = "",
            clean: bool = True) -> dict:
    """Full pipeline.

    Args:
        config:      Orbi config (used for the LLM client)
        source_path: file inside the workspace folder
        target:     'pdf' / 'docx' / 'txt' / 'md' / 'xlsx' / 'csv'
        out_dir:    where to save the result (typically the workspace folder)
        hint:       optional owner instruction ("make it short", "use bullets", etc.)
        clean:      if False, skip the LLM cleanup step (raw conversion)

    Returns: {output_path, output_name, original_chars, cleaned_chars, skipped_clean}
    """
    target = target.lower().strip()
    if target not in SUPPORTED_TARGETS:
        raise ValueError(f"unsupported target {target!r}; valid: {SUPPORTED_TARGETS}")
    if not source_path.exists():
        raise FileNotFoundError(str(source_path))

    raw = extract_text(source_path)
    if not raw or not raw.strip():
        raise ValueError("could not extract any text from the source file "
                         "(maybe it's a scanned PDF or image — those need OCR, "
                         "which is a future feature)")

    kind = detect_kind(source_path)
    cleaned = llm_clean(config, raw, target, kind=kind, hint=hint) if clean else raw

    out_dir.mkdir(parents=True, exist_ok=True)
    stem = source_path.stem
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    suffix = "." + target
    out_name = f"{stem}__cleaned_{timestamp}{suffix}"
    out_path = out_dir / out_name

    title = f"Cleaned: {source_path.name}"
    if target == "pdf":
        write_pdf(cleaned, out_path, title=title)
    elif target == "docx":
        write_docx(cleaned, out_path, title=title)
    elif target == "txt":
        write_txt(cleaned, out_path)
    elif target == "md":
        write_md(cleaned, out_path)
    elif target == "xlsx":
        write_xlsx(cleaned, out_path, title=title)
    elif target == "csv":
        write_csv(cleaned, out_path)
    else:
        raise ValueError(f"unhandled target {target}")

    return {
        "output_path":    str(out_path),
        "output_name":    out_name,
        "original_chars": len(raw),
        "cleaned_chars":  len(cleaned),
        "skipped_clean":  not clean,
        "source":         source_path.name,
        "target":         target,
    }
