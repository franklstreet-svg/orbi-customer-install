"""
modules/form_filler — fill a labeled template with real project/CO data.

Given:
  - a template record from modules/forms.py (which knows the file path,
    its format, and the field → Orby-data-key mapping)
  - a context dict of Orby data values (customer_name, address, amount, ...)

Produce a filled output file (PDF or DOCX) and return the path to it.

PDF path: uses pypdf to fill AcroForm fields in place.
DOCX path: uses python-docx to substitute {{var}} / [Var Name] / <<VAR>>
placeholders in paragraphs + table cells.
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path


# Build a context dict from a project + (optional) change-order + business.
# Returns the canonical-keyed dict that templates.detected_fields[].mapped_to
# can look values up in.
def build_context_for_co(project: dict, co: dict | None,
                          business: dict) -> dict:
    """Standard context for change-order forms. Pulls customer/project
    data, the CO amount/description if provided, and the business identity."""
    contract = float(project.get("contract_amount") or 0)
    co_amount = float((co or {}).get("amount") or 0)
    new_total = contract + co_amount
    today_str = datetime.now().strftime("%B %-d, %Y")

    biz_addr = business.get("address") or {}
    if isinstance(biz_addr, dict):
        biz_addr_str = ", ".join(filter(None, [
            (biz_addr.get("street") or "").strip(),
            (biz_addr.get("city") or "").strip(),
            f"{biz_addr.get('state') or ''} {biz_addr.get('zip') or ''}".strip(),
        ]))
    else:
        biz_addr_str = str(biz_addr)

    contact = business.get("contact") or {}

    # Project address may itself be a flat string
    project_addr = project.get("address") or ""

    return {
        "customer_name":    project.get("customer_name") or "",
        "customer_phone":   project.get("customer_phone") or "",
        "customer_email":   project.get("customer_email") or "",
        "address":          project_addr,
        "city":             "",
        "state":            "",
        "zip":              "",
        "project_label":    project.get("label") or "",
        "co_number":        f"#{(co or {}).get('co_number','?')}" if co else "",
        "contract_amount":  f"${contract:,.2f}",
        "amount":           f"${co_amount:,.2f}" if co else "",
        "new_contract_amount": f"${new_total:,.2f}" if co else f"${contract:,.2f}",
        "description":      (co or {}).get("description") or "",
        "today":            today_str,
        "contracted_at":    project.get("contracted_at") or "",
        "started_at":       project.get("started_at") or "",
        "est_complete":     project.get("est_complete") or "",
        "biz_name":         business.get("name") or "",
        "biz_license":      business.get("license") or "",
        "biz_address":      biz_addr_str,
        "biz_phone":        contact.get("phone") or "",
        "biz_email":        contact.get("email") or "",
        # Signature fields stay blank — filled by the customer at sign time
        "signature_pad":    "",
        "biz_signature_pad":"",
    }


def fill(template_path: Path, fmt: str,
          detected_fields: list[dict],
          context: dict,
          out_path: Path) -> Path:
    """Top-level fill. Dispatches by format. Returns out_path on success."""
    template_path = Path(template_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "pdf":
        return _fill_pdf(template_path, detected_fields, context, out_path)
    if fmt == "docx":
        return _fill_docx(template_path, detected_fields, context, out_path)
    raise ValueError(f"unsupported format: {fmt}")


# ── PDF AcroForm fill ────────────────────────────────────────────────────


def _fill_pdf(template_path: Path, detected_fields: list[dict],
               context: dict, out_path: Path) -> Path:
    import pypdf
    # Build {pdf_field_name: value} dict
    values: dict[str, str] = {}
    for f in detected_fields or []:
        pdf_name = f.get("name")
        mapped = (f.get("mapped_to") or "").strip()
        if not pdf_name or not mapped:
            continue
        val = context.get(mapped)
        if val is None or val == "":
            continue
        values[pdf_name] = str(val)

    reader = pypdf.PdfReader(str(template_path))
    writer = pypdf.PdfWriter()
    # Copy every page over
    for page in reader.pages:
        writer.add_page(page)
    # update_page_form_field_values writes the values into the form
    if values:
        try:
            for page in writer.pages:
                writer.update_page_form_field_values(page, values)
        except Exception:
            # Older pypdf signature differences — fall back to per-page calls
            for page in writer.pages:
                try:
                    writer.update_page_form_field_values(page, values)
                except Exception:
                    continue
    with out_path.open("wb") as f:
        writer.write(f)
    return out_path


# ── DOCX placeholder substitution ───────────────────────────────────────


# Same patterns the detector found: {{x}} / [X] / <<x>>
_PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([A-Za-z][\w\s]+?)\s*\}\}"
    r"|\[\s*([A-Za-z][\w\s]+?)\s*\]"
    r"|<<\s*([A-Za-z][\w\s]+?)\s*>>"
)


def _fill_docx(template_path: Path, detected_fields: list[dict],
                context: dict, out_path: Path) -> Path:
    import docx
    # name → mapped_to lookup (for use in the regex callback)
    name_to_mapped = {(f.get("name") or "").strip(): (f.get("mapped_to") or "").strip()
                       for f in (detected_fields or [])}

    def _replace(text: str) -> str:
        def _sub(m):
            name = (m.group(1) or m.group(2) or m.group(3) or "").strip()
            mapped = name_to_mapped.get(name) or name_to_mapped.get(name.lower()) or ""
            if not mapped:
                return m.group(0)  # leave placeholder alone if unmapped
            val = context.get(mapped)
            return "" if val is None else str(val)
        return _PLACEHOLDER_RE.sub(_sub, text)

    d = docx.Document(str(template_path))
    for p in d.paragraphs:
        new_text = _replace(p.text)
        if new_text != p.text:
            # Clear existing runs and set the replaced text on a single run
            # Note: this loses character-level formatting within the
            # replaced span. Trade-off for v1.
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = new_text
            else:
                p.add_run(new_text)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    new_text = _replace(p.text)
                    if new_text != p.text:
                        for run in p.runs:
                            run.text = ""
                        if p.runs:
                            p.runs[0].text = new_text
                        else:
                            p.add_run(new_text)
    d.save(str(out_path))
    return out_path
