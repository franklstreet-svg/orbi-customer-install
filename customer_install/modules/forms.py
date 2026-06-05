"""
modules/forms — form template registry.

The contractor uploads their own change-order, contract, lien-waiver, etc.
templates. Orby reads each one, detects fillable fields (PDF AcroForm or
DOCX placeholders), maps each detected field to one of Orby's known data
fields, and remembers the mapping forever for that template.

When a chat handler later asks Orby to "fill out the change order for the
Johnson project," she:
  1. picks the latest template of kind="change_order"
  2. looks up the field-mapping
  3. builds a context dict (customer_name → project.customer_name, etc.)
  4. hands it to modules/form_filler.py to produce the filled output

Template record shape:
  {
    "id":           "12-char hex",
    "kind":         "change_order"|"contract"|"lien_waiver"|"w9"|
                     "coi_request"|"subcontractor_agreement"|"punch_list"|
                     "custom",
    "display_name": "Acme Standard CO Template",
    "filename":     "Acme_Construction_Change_Order.pdf",
    "storage_path": "data/form_templates/<id>.pdf",
    "format":       "pdf" | "docx",
    "detected_fields": [
       {"name": "Customer Name", "type": "text", "mapped_to": "customer_name"},
       {"name": "Project Site",  "type": "text", "mapped_to": "address"},
       ...
    ],
    "uploaded_at": 1780000000,
    "uploaded_by": "frank",
    "version":      1,                    # bumped on re-upload of same kind
    "is_default":   true,                 # default for its kind
  }

KEY DESIGN: there's ONE default template per kind. When a chat handler
needs "the change order template," it loads the default for kind="change_order".
The owner can upload multiple templates of the same kind (e.g. residential
vs commercial CO) and switch the default.
"""
from __future__ import annotations

import json
import re
import threading
import time
import uuid
from pathlib import Path

_LOCK = threading.Lock()
FILE = "form_templates.json"
TEMPLATES_DIR = "form_templates"

# Recognized kinds. Owner picks one when uploading.
VALID_KINDS = {
    "change_order",
    "contract",
    "msa",                       # master service agreement
    "lien_waiver_partial_cond",
    "lien_waiver_partial_uncond",
    "lien_waiver_final_cond",
    "lien_waiver_final_uncond",
    "w9",
    "coi_request",
    "subcontractor_agreement",
    "punch_list",
    "proposal",
    "invoice_custom",
    "custom",
}

# ── Field name → Orby data field mapping (synonym dictionary) ─────────────
# The auto-mapper's first pass. Lowercase + strip punctuation before lookup.
# Each entry: list of common variants → canonical Orby data-field key.
_FIELD_SYNONYMS: list[tuple[list[str], str]] = [
    # Customer identity. Dropped bare "name" — too greedy (it matches
    # "biz_name", "project_name", etc. via loose-match).
    (["customer name", "customer", "client name", "client", "owner name",
      "owner", "homeowner", "homeowner name", "buyer", "buyer name",
      "mr mrs ms", "your name", "purchaser", "customer fullname",
      "client fullname"],
     "customer_name"),
    (["customer phone", "phone", "phone number", "client phone",
      "owner phone", "contact phone", "tel", "telephone"],
     "customer_phone"),
    (["customer email", "email", "email address", "client email",
      "owner email", "contact email"],
     "customer_email"),
    # Project location
    (["job address", "project address", "project site", "site address",
      "property address", "property", "location", "address", "site",
      "work address", "project location", "job site"],
     "address"),
    (["city"], "city"),
    (["state"], "state"),
    (["zip", "zip code", "postal code"], "zip"),
    # Project + CO meta
    (["job name", "project name", "project", "label"], "project_label"),
    (["co number", "co no", "co #", "change order number",
      "change order #", "co num", "number", "co"],
     "co_number"),
    (["original contract", "contract amount", "contract price",
      "base contract", "original contract amount", "contract total",
      "contract value", "total contract"],
     "contract_amount"),
    (["co amount", "change amount", "amount of change",
      "amount of this change order", "this change", "amount",
      "change order amount", "additional amount", "extra"],
     "amount"),
    (["new contract", "new contract amount", "revised contract",
      "revised total", "updated contract", "new total"],
     "new_contract_amount"),
    (["description", "scope", "work to be done", "scope of work",
      "details", "work description", "description of work",
      "description of change", "what changed"],
     "description"),
    # Dates
    (["date", "today", "today's date", "today date", "co date",
      "issue date", "date of change", "date issued"],
     "today"),
    (["contract date", "original contract date"], "contracted_at"),
    (["start date", "started", "project start"], "started_at"),
    (["completion date", "est complete", "estimated completion",
      "expected completion"], "est_complete"),
    # Contractor identity (from business_info)
    (["contractor", "contractor name", "company", "company name",
      "from", "general contractor", "gc",
      "biz", "biz name", "business name", "business",
      "contractor company", "your company"],
     "biz_name"),
    (["contractor license", "license", "license number", "license #",
      "license no", "state license", "contractor license number",
      "contractor license #"],
     "biz_license"),
    (["contractor address", "company address", "office address"],
     "biz_address"),
    (["contractor phone", "company phone", "office phone"],
     "biz_phone"),
    (["contractor email", "company email", "office email"],
     "biz_email"),
    # Signature
    (["customer signature", "client signature", "owner signature",
      "homeowner signature", "buyer signature", "purchaser signature",
      "owner sign here", "client sign", "sign here", "signature"],
     "signature_pad"),
    (["contractor signature", "gc signature", "company signature"],
     "biz_signature_pad"),
]


def _path(data_dir: Path) -> Path:
    return Path(data_dir) / FILE


def _templates_dir(data_dir: Path) -> Path:
    p = Path(data_dir) / TEMPLATES_DIR
    p.mkdir(parents=True, exist_ok=True)
    return p


def _load(data_dir: Path) -> list[dict]:
    p = _path(data_dir)
    if not p.exists():
        return []
    try:
        return json.loads(p.read_text(encoding="utf-8") or "[]")
    except (json.JSONDecodeError, OSError):
        return []


def _save(data_dir: Path, templates: list[dict]) -> None:
    p = _path(data_dir)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(templates, indent=2), encoding="utf-8")
    tmp.replace(p)


# ── Public API ───────────────────────────────────────────────────────────


def add(data_dir: Path, *,
        kind: str, display_name: str, filename: str,
        source_path: Path, uploaded_by: str = "owner",
        is_default: bool = True) -> dict:
    """Register a new template. Copies the file into data/form_templates/
    under a unique name, detects fields, auto-maps them, and persists."""
    if kind not in VALID_KINDS:
        raise ValueError(f"unknown kind {kind!r}; expected one of {VALID_KINDS}")
    src = Path(source_path)
    if not src.exists():
        raise FileNotFoundError(f"source file missing: {src}")
    suffix = src.suffix.lower()
    if suffix not in (".pdf", ".docx"):
        raise ValueError(f"only .pdf or .docx supported (got {suffix})")
    tpl_id = uuid.uuid4().hex[:12]
    fmt = "pdf" if suffix == ".pdf" else "docx"
    stored = _templates_dir(data_dir) / f"{tpl_id}{suffix}"
    stored.write_bytes(src.read_bytes())

    # Detect fields and auto-map
    detected = detect_fields(stored, fmt)
    for f in detected:
        f["mapped_to"] = auto_map_field_name(f["name"]) or ""

    now = int(time.time())
    with _LOCK:
        templates = _load(data_dir)
        # If marking default, demote any existing default of the same kind
        if is_default:
            for t in templates:
                if t.get("kind") == kind and t.get("is_default"):
                    t["is_default"] = False
        # Version: max(existing version for this kind) + 1
        existing_versions = [int(t.get("version", 1)) for t in templates
                              if t.get("kind") == kind]
        version = (max(existing_versions) + 1) if existing_versions else 1
        record = {
            "id":              tpl_id,
            "kind":            kind,
            "display_name":    (display_name or "").strip() or filename,
            "filename":        filename,
            "storage_path":    str(stored.relative_to(Path(data_dir).parent)) if Path(data_dir).is_absolute() else str(stored),
            "format":          fmt,
            "detected_fields": detected,
            "uploaded_at":     now,
            "uploaded_by":     uploaded_by,
            "version":         version,
            "is_default":      bool(is_default),
        }
        templates.append(record)
        _save(data_dir, templates)
    return record


def list_all(data_dir: Path, kind: str | None = None) -> list[dict]:
    templates = _load(data_dir)
    if kind:
        templates = [t for t in templates if t.get("kind") == kind]
    # Sort: default first, then newest first
    templates.sort(key=lambda t: (not t.get("is_default", False),
                                    -int(t.get("uploaded_at", 0))))
    return templates


def get(data_dir: Path, template_id: str) -> dict | None:
    for t in _load(data_dir):
        if t.get("id") == template_id:
            return t
    return None


def get_default(data_dir: Path, kind: str) -> dict | None:
    """The default template for a given kind. Returns None if none uploaded.
    Used by chat handlers — 'fill out the change order' → get_default('change_order')."""
    for t in list_all(data_dir, kind=kind):
        if t.get("is_default"):
            return t
    # No explicit default — return the newest of this kind
    candidates = list_all(data_dir, kind=kind)
    return candidates[0] if candidates else None


def set_default(data_dir: Path, template_id: str) -> bool:
    with _LOCK:
        templates = _load(data_dir)
        target = next((t for t in templates if t.get("id") == template_id), None)
        if not target:
            return False
        kind = target.get("kind")
        for t in templates:
            t["is_default"] = (t.get("id") == template_id) if t.get("kind") == kind else t.get("is_default", False)
        _save(data_dir, templates)
        return True


def update_field_mapping(data_dir: Path, template_id: str,
                          field_name: str, mapped_to: str) -> bool:
    """Owner-side correction of an auto-mapping. Persists to disk."""
    with _LOCK:
        templates = _load(data_dir)
        for t in templates:
            if t.get("id") != template_id:
                continue
            for f in t.get("detected_fields", []) or []:
                if f.get("name") == field_name:
                    f["mapped_to"] = (mapped_to or "").strip()
                    _save(data_dir, templates)
                    return True
    return False


def remove(data_dir: Path, template_id: str) -> bool:
    with _LOCK:
        templates = _load(data_dir)
        before = len(templates)
        target = next((t for t in templates if t.get("id") == template_id), None)
        if not target:
            return False
        templates = [t for t in templates if t.get("id") != template_id]
        _save(data_dir, templates)
        # Also delete the stored file
        try:
            stored = _templates_dir(data_dir) / f"{template_id}.{target.get('format','pdf')}"
            if stored.exists():
                stored.unlink()
        except OSError:
            pass
        return len(templates) < before


# ── Detection ────────────────────────────────────────────────────────────


def detect_fields(template_path: Path, fmt: str) -> list[dict]:
    """Enumerate fillable fields in the template. Returns
    [{"name": "Customer Name", "type": "text"|"checkbox"|"signature"}, ...].
    Empty list if no fillable fields found (template is a flat scan or
    has no placeholders)."""
    template_path = Path(template_path)
    if fmt == "pdf":
        return _detect_pdf_fields(template_path)
    if fmt == "docx":
        return _detect_docx_placeholders(template_path)
    return []


def _detect_pdf_fields(path: Path) -> list[dict]:
    try:
        import pypdf
    except ImportError:
        return []
    out = []
    try:
        with path.open("rb") as f:
            reader = pypdf.PdfReader(f)
            fields = reader.get_fields() or {}
            for name, fld in fields.items():
                # pypdf may give us a dict-like or a PdfObject
                ft = ""
                try:
                    ft = (fld.get("/FT") or "").strip("/").lower() if hasattr(fld, "get") else ""
                except Exception:
                    pass
                kind = {
                    "tx": "text",
                    "btn": "checkbox",
                    "ch": "choice",
                    "sig": "signature",
                }.get(ft, "text")
                out.append({"name": name, "type": kind})
    except Exception:
        return []
    return out


# Word placeholder patterns: {{customer_name}}, [Customer Name], <<CUSTOMER>>
_DOCX_PLACEHOLDER_RE = re.compile(
    r"(?:\{\{\s*([A-Za-z][\w\s]+?)\s*\}\}"
    r"|\[\s*([A-Za-z][\w\s]+?)\s*\]"
    r"|<<\s*([A-Za-z][\w\s]+?)\s*>>)"
)


def _detect_docx_placeholders(path: Path) -> list[dict]:
    try:
        import docx
    except ImportError:
        return []
    try:
        d = docx.Document(str(path))
    except Exception:
        return []
    found: dict[str, dict] = {}
    for p in d.paragraphs:
        for m in _DOCX_PLACEHOLDER_RE.finditer(p.text):
            name = (m.group(1) or m.group(2) or m.group(3) or "").strip()
            if name and name not in found:
                found[name] = {"name": name, "type": "text"}
    # Also walk tables
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for m in _DOCX_PLACEHOLDER_RE.finditer(cell.text):
                    name = (m.group(1) or m.group(2) or m.group(3) or "").strip()
                    if name and name not in found:
                        found[name] = {"name": name, "type": "text"}
    return list(found.values())


# ── Auto-mapping ─────────────────────────────────────────────────────────


def _normalize_field_name(name: str) -> str:
    """Strip punctuation, collapse whitespace, lowercase. Used by the
    synonym-dictionary lookup. Also flattens underscores + hyphens to
    spaces so {{contract_amount}} and 'Contract Amount' normalize the same."""
    s = (name or "").lower()
    # Underscores and hyphens become spaces (so "contract_amount" → "contract amount")
    s = re.sub(r"[_\-]+", " ", s)
    # Other non-word/non-space punctuation → space
    s = re.sub(r"[^\w\s]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def auto_map_field_name(field_name: str) -> str | None:
    """Pass 1 of the auto-mapper: synonym-dictionary lookup. Returns the
    canonical Orby data-field key or None if no match. Owner can override
    via update_field_mapping()."""
    normalized = _normalize_field_name(field_name)
    if not normalized:
        return None
    # Exact normalized match against the synonym table
    for synonyms, canonical in _FIELD_SYNONYMS:
        if normalized in synonyms:
            return canonical
    # Loose match: contains a synonym as substring (only for long names)
    for synonyms, canonical in _FIELD_SYNONYMS:
        for s in synonyms:
            if len(s) >= 4 and s in normalized:
                return canonical
    return None


def summary(data_dir: Path) -> dict:
    """For the dashboard Forms tab — totals + per-kind counts."""
    templates = _load(data_dir)
    by_kind: dict[str, int] = {}
    for t in templates:
        k = t.get("kind", "?")
        by_kind[k] = by_kind.get(k, 0) + 1
    return {
        "total":   len(templates),
        "by_kind": by_kind,
        "kinds_with_default": [
            k for k in by_kind
            if any(t.get("is_default") for t in templates if t.get("kind") == k)
        ],
    }
